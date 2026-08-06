"""Render the same report content as a Word document.

Chart-first design mirrors html.py: each chart gets its own card (indicator
name, a one-line plain-language definition, a scope line, the chart, one
caption). The Analyst/Anomaly Reviewer's fuller narratives move to a final
"Detailed notes" appendix page rather than the main body.
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

OUTPUT_DIR = Path(__file__).resolve().parent.parent.parent / "output"

STATUS_LABEL = {"green": "On track", "amber": "Worth a look", "red": "Needs attention"}
STATUS_COLOR = {
    "green": RGBColor(0x0C, 0xA3, 0x0C),
    "amber": RGBColor(0xB8, 0x79, 0x0F),
    "red": RGBColor(0xD0, 0x3B, 0x3B),
}
INK_SECONDARY = RGBColor(0x52, 0x51, 0x4E)
INK_MUTED = RGBColor(0x89, 0x87, 0x81)


def _status_line(doc: Document, status: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"Status: {STATUS_LABEL[status]}")
    run.bold = True
    run.font.color.rgb = STATUS_COLOR[status]
    run.font.size = Pt(10)


def _caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = INK_SECONDARY


def _fmt_trend(trend: dict) -> str:
    if not trend:
        return "Not enough data"
    arrow = {"up": "up", "down": "down", "flat": "flat"}[trend["direction"]]
    pct = f"{trend['pct_change']:+.1f}%" if trend["pct_change"] is not None else "n/a"
    return f"{arrow} {pct} ({trend['start_year']}-{trend['end_year']})"


def _fmt_anomaly_change(a: dict) -> str:
    """Compact "Change" column for the anomaly table: a % for a year-over-year
    flag, a standard-deviation count for a baseline-deviation flag - the two
    anomaly checks in src/analysis/anomalies.py measure different things, so
    they get different units rather than a misleading shared one."""
    if a["type"] == "yoy_change":
        return f"{a['metric']:+.1f}%"
    return f"{a['metric']:+.1f} SD vs baseline"


def _scope_line(report: dict, first_year: int, last_year: int, note: str = None) -> str:
    base = f"Saudi Arabia ({report['country_display']}), {first_year} - {last_year}"
    return f"{base} ({note})" if note else base


def _comparison_note(has_region: bool, has_global: bool):
    if has_region and has_global:
        return None
    missing = []
    if not has_region:
        missing.append("regional")
    if not has_global:
        missing.append("global")
    return f"{' and '.join(missing)} comparison data not available for this indicator"


def _add_stats_table(doc: Document, stats: dict) -> None:
    latest = stats["latest"]
    table = doc.add_table(rows=2, cols=3)
    table.style = "Light Grid Accent 1"
    headers = [f"Latest value ({latest['year']})", "5-year change", "10-year change"]
    values = [str(latest["value"]), _fmt_trend(stats.get("trend_5y")), _fmt_trend(stats.get("trend_10y"))]
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.paragraphs[0].add_run(text).bold = True
    for i, text in enumerate(values):
        table.rows[1].cells[i].paragraphs[0].add_run(text)


def _add_anomaly_table(doc: Document, anomalies: list) -> None:
    if not anomalies:
        doc.add_paragraph("No anomalies detected.")
        return
    table = doc.add_table(rows=1 + len(anomalies), cols=3)
    table.style = "Light Grid Accent 1"
    for i, header in enumerate(["Year", "Change", "Severity"]):
        table.rows[0].cells[i].paragraphs[0].add_run(header).bold = True
    for r, a in enumerate(anomalies, start=1):
        row = table.rows[r]
        row.cells[0].text = str(a["year"])
        row.cells[1].text = _fmt_anomaly_change(a)
        severity_run = row.cells[2].paragraphs[0].add_run(a["severity"].title())
        if a["severity"] in ("high", "medium"):
            severity_run.font.color.rgb = STATUS_COLOR["red" if a["severity"] == "high" else "amber"]
            severity_run.bold = True


def _add_chart_card(doc: Document, title: str, definition: str, scope: str, chart_path: Path, caption: str) -> None:
    doc.add_heading(title, level=2)

    def_p = doc.add_paragraph()
    def_p.add_run(definition).font.size = Pt(11)

    scope_p = doc.add_paragraph()
    scope_run = scope_p.add_run(scope.upper())
    scope_run.font.size = Pt(9)
    scope_run.font.color.rgb = INK_MUTED

    doc.add_picture(str(chart_path), width=Inches(6))
    _caption(doc, caption)


def render_docx(
    report: dict,
    narrative: dict,
    chart_paths: dict,
    trend_status: str,
    anomaly_status: str,
    overall_status: str,
    has_region: bool,
    has_global: bool,
) -> Document:
    doc = Document()
    stats = report["stats"]
    yoy_years = [p["year"] for p in stats["yoy_changes"]]
    first_year, last_year = yoy_years[0], yoy_years[-1]
    indicator_name = report["indicator_name"]

    title = doc.add_heading(f"{report['disease']} in {report['country_display']}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    subtitle = doc.add_paragraph()
    run = subtitle.add_run(f"{report['indicator_name']} · generated {report['generated_date']}")
    run.font.color.rgb = INK_MUTED
    run.font.size = Pt(10)

    doc.add_heading("Executive summary", level=1)
    _status_line(doc, overall_status)
    doc.add_paragraph(narrative["executive_summary"])
    _add_stats_table(doc, stats)

    _add_chart_card(
        doc,
        indicator_name,
        narrative["definition"],
        _scope_line(report, first_year, last_year),
        chart_paths["line_chart"],
        narrative["trend_caption"],
    )
    _add_chart_card(
        doc,
        indicator_name,
        narrative["definition"],
        _scope_line(report, first_year, last_year),
        chart_paths["bar_chart"],
        narrative["change_caption"],
    )
    _add_chart_card(
        doc,
        indicator_name,
        narrative["definition"],
        _scope_line(report, first_year, last_year, _comparison_note(has_region, has_global)),
        chart_paths["comparison_chart"],
        narrative["comparison_caption"],
    )

    doc.add_heading("Anomalies", level=1)
    _status_line(doc, anomaly_status)
    _add_anomaly_table(doc, report["anomalies"])

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer_run = footer.add_run(
        f"Source: WHO Global Health Observatory (indicator {report['indicator_code']}). "
        "This report is generated automatically from statistical analysis and is not medical advice."
    )
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = INK_MUTED

    doc.add_page_break()
    doc.add_heading("Detailed notes", level=1)
    doc.add_heading("Analyst interpretation", level=2)
    doc.add_paragraph(narrative["analyst"])
    doc.add_heading("Anomaly review", level=2)
    doc.add_paragraph(narrative["anomaly_reviewer"])

    return doc


def save_docx(
    report: dict,
    narrative: dict,
    chart_paths: dict,
    trend_status: str,
    anomaly_status: str,
    overall_status: str,
    has_region: bool,
    has_global: bool,
) -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    stem = f"{report['indicator_code']}_{report['country_display']}_{report['generated_date']}"
    path = OUTPUT_DIR / f"{stem}.docx"
    doc = render_docx(
        report, narrative, chart_paths, trend_status, anomaly_status, overall_status, has_region, has_global
    )
    doc.save(str(path))
    return path
